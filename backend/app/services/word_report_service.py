"""Generate the business plan report as a styled .docx (python-docx)."""
from __future__ import annotations

import io
import re
import uuid
from datetime import datetime, timezone
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from ..models import BusinessPlanProject
from ..schemas.reports import ReportFile
from . import chart_render_service as charts
from . import report_data_service as rd

NAVY = RGBColor(0x1E, 0x29, 0x3B)
BLUE = RGBColor(0x25, 0x63, 0xEB)
SLATE = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1E293B"
SUBTOTAL_BG = "F1F5F9"
SECTION_BG = "E2E8F0"
AMBER_BG = "FEF3C7"
RED_BG = "FEE2E2"
BLUE_BG = "EFF6FF"


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------
def _shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _no_borders_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "CBD5E1")
        borders.append(e)
    tblPr.append(borders)


def _page_number_field(paragraph):
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        run = paragraph.add_run()
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x33, 0x3D, 0x4F)
    for i, size in ((1, 16), (2, 13), (3, 11)):
        st = doc.styles[f"Heading {i}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY if i == 1 else BLUE
        st.font.bold = True


def _start_section(doc, landscape: bool):
    """Begin a new page-section, switching orientation. New sections inherit the
    running header/footer from the previous one (linked by default)."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    w, h = section.page_width, section.page_height
    long_side, short_side = max(w, h), min(w, h)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = long_side, short_side
        section.left_margin = section.right_margin = Inches(0.6)
        section.top_margin = section.bottom_margin = Inches(0.7)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = short_side, long_side
        section.left_margin = section.right_margin = Inches(0.8)
        section.top_margin = section.bottom_margin = Inches(0.9)
    return section


def add_heading(doc, text, level=1, page_break=True):
    if level == 1 and page_break:
        add_page_break(doc)
    h = doc.add_heading(text, level=level)
    if level == 1:
        p = h._p
        pPr = p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "2563EB")
        bottom.set(qn("w:space"), "4")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return h


def add_subheading(doc, text):
    return doc.add_heading(text, level=3)


def add_paragraph_block(doc, text, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_page_break(doc):
    doc.add_page_break()


# --------------------------------------------------------------------------
# tables
# --------------------------------------------------------------------------
def add_financial_table(doc, periods, rows, currency, total_col=False):
    cols = 1 + len(periods) + (1 if total_col else 0)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders_table(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "Line item", left=True)
    for i, p in enumerate(periods):
        _hdr_cell(hdr[1 + i], p)
    if total_col:
        _hdr_cell(hdr[-1], "Total")

    for r in rows:
        cells = table.add_row().cells
        kind = r.get("kind", "line")
        label = r["label"]
        bold = kind in ("subtotal", "grand", "section", "group")
        _label_cell(cells[0], label, bold=bold, indent=(kind == "line"))
        vals = r.get("values", [])
        for i in range(len(periods)):
            v = vals[i] if i < len(vals) else None
            _num_cell(cells[1 + i], rd.fmt_num(v) if (kind != "section" and vals) else "", bold=bold)
        if total_col:
            _num_cell(cells[-1], rd.fmt_num(sum(vals)) if vals and kind != "section" else "", bold=bold)
        if kind in ("subtotal", "grand", "check"):
            for c in cells:
                _shade(c, SUBTOTAL_BG)
        elif kind in ("section", "group"):
            for c in cells:
                _shade(c, SECTION_BG)
    return table


def add_kpi_table(doc, items, currency=None, cols=2):
    """items: list of {label, value} -> a 2-column key/value grid."""
    pairs = items if isinstance(items, list) else [{"label": k, "value": v} for k, v in items.items()]
    table = doc.add_table(rows=0, cols=2)
    _no_borders_table(table)
    for it in pairs:
        cells = table.add_row().cells
        _label_cell(cells[0], it["label"], bold=False)
        val = it["value"]
        _num_cell(cells[1], rd.fmt_num(val) if isinstance(val, (int, float)) else str(val), bold=True)
    return table


def add_generic_table(doc, headers, rows, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders_table(table)
    for i, h in enumerate(headers):
        _hdr_cell(table.rows[0].cells[i], h, left=(i not in numeric_cols))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i in numeric_cols:
                _num_cell(cells[i], rd.fmt_num(val) if isinstance(val, (int, float)) else str(val))
            else:
                _label_cell(cells[i], str(val))
    return table


def add_warning_box(doc, message, severity):
    bg = {"critical": RED_BG, "warning": AMBER_BG}.get(severity, BLUE_BG)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(f"[{severity.upper()}] ")
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(message).font.size = Pt(9)


def _hdr_cell(cell, text, left=False):
    _shade(cell, HEADER_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE


def _label_cell(cell, text, bold=False, indent=False):
    p = cell.paragraphs[0]
    run = p.add_run(("   " if indent else "") + str(text))
    run.bold = bold
    run.font.size = Pt(9)


def _num_cell(cell, text, bold=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)


# --------------------------------------------------------------------------
# header / footer / cover
# --------------------------------------------------------------------------
def _header_footer(doc, ctx):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{ctx['meta']['company']}  |  Business Plan Financial Projection"
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.color.rgb = SLATE

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"Confidential Business Plan Report — Generated by {ctx['app_name']}   ·   Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = SLATE
    _page_number_field(fp)


def add_cover_page(doc, ctx):
    m = ctx["meta"]
    for _ in range(3):
        doc.add_paragraph()
    def big(text, size, color, bold=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p
    big(m["company"], 26, NAVY)
    if m["project_name"] and m["project_name"] != "–":
        big(m["project_name"], 13, SLATE, bold=False)
    doc.add_paragraph()
    big(m["title"], 18, BLUE)
    big(m["subtitle"], 12, SLATE, bold=False)
    doc.add_paragraph()
    for line in (f"Currency: {m['currency']}", f"Scenario: {m['scenario_label']}",
                 f"Projection period: {m['period_range']}", f"Prepared for: {m['prepared_for']}",
                 f"Prepared: {m['prepared_date']}"):
        big(line, 11, SLATE, bold=False)
    for _ in range(2):
        doc.add_paragraph()
    c = big("CONFIDENTIAL", 11, RGBColor(0xB9, 0x1C, 0x1C))
    c.runs[0].bold = True


def add_table_of_contents_placeholder(doc):
    add_heading(doc, "Table of Contents", level=1)
    add_paragraph_block(doc, "Right-click the list below and choose 'Update Field' in Microsoft Word to refresh page numbers.", italic=True, size=9)
    for i, (_k, title) in enumerate(rd.SECTIONS, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}.  {title}").font.size = Pt(10)


# --------------------------------------------------------------------------
# sections
# --------------------------------------------------------------------------
def add_executive_summary(doc, ctx):
    add_heading(doc, "1. Executive Summary", level=1)
    ov = ctx["overview"]
    desc = ov.get("Description", "")
    add_paragraph_block(doc, desc)
    add_paragraph_block(doc,
        f"This report presents a {ctx['meta']['period_range']} projected financial study for "
        f"{ctx['meta']['company']} under the {ctx['scenario_label']}, prepared in {ctx['currency']}.")
    add_subheading(doc, "Key projected highlights")
    add_kpi_table(doc, ctx["highlights"])
    if ctx["insights"]:
        add_subheading(doc, "Analytical commentary")
        for s in ctx["insights"]:
            doc.add_paragraph(s, style="List Bullet")


def add_business_overview(doc, ctx):
    add_heading(doc, "2. Business Overview", level=1)
    add_kpi_table(doc, ctx["overview"])


def add_products_revenue(doc, ctx):
    add_heading(doc, "3. Products and Revenue Model", level=1)
    add_subheading(doc, "Products and services")
    rows = [[p["name"], p["category"], p["revenue_type"], p["unit"], p["price"], p["launch"], p["active"]]
            for p in ctx["products"]]
    add_generic_table(doc, ["Product / Service", "Category", "Revenue type", "Unit", "Price", "Launch", "Status"],
                      rows, numeric_cols=(4,))
    add_subheading(doc, "Revenue by stream")
    add_financial_table(doc, ctx["periods"], _cat_rows(ctx["revenue_table"], "Total revenue"), ctx["currency"], total_col=True)


def add_assumptions_sections(doc, ctx):
    if not ctx["options"].include_assumptions:
        return
    add_heading(doc, "4. Projection Assumptions", level=1)
    add_subheading(doc, "Direct cost of sales")
    add_financial_table(doc, ctx["periods"], _cat_rows(ctx["direct_cost_table"], "Total cost of sales"), ctx["currency"], total_col=True)
    add_subheading(doc, "Operating expenses")
    add_financial_table(doc, ctx["periods"], _cat_rows(ctx["opex_table"], "Total operating expenses"), ctx["currency"], total_col=True)

    add_heading(doc, "5. Operating Model", level=1)
    add_subheading(doc, "Staffing plan")
    rows = [[s["department"], s["title"], s["employees"], s["salary"], s["start"]] for s in ctx["staffing"]]
    add_generic_table(doc, ["Department", "Role", "Employees", "Monthly salary", "Start date"], rows, numeric_cols=(2, 3))
    if ctx["staff_by_dept"]:
        add_subheading(doc, "Total staff cost by year")
        add_financial_table(doc, ctx["periods"], _cat_rows(ctx["staff_by_dept"], "Total staff cost"), ctx["currency"], total_col=True)

    add_heading(doc, "6. Capital Expenditure and Assets", level=1)
    rows = [[a["name"], a["category"], a["purchase_amount"], a["useful_life"], a["method"], a["nbv"]] for a in ctx["fixed_assets"]]
    add_generic_table(doc, ["Asset", "Category", "Cost", "Life (yrs)", "Method", "NBV (final year)"], rows, numeric_cols=(2, 3, 5))

    add_heading(doc, "7. Working Capital and Financing", level=1)
    add_subheading(doc, "Working capital assumptions")
    add_kpi_table(doc, ctx["working_capital"])
    add_subheading(doc, "Equity funding")
    add_kpi_table(doc, ctx["financing"]["equity"])
    if ctx["financing"]["loans"]:
        add_subheading(doc, "Loans")
        rows = [[l["name"], l["amount"], rd.fmt_pct(l["rate"]), l["term"], l["grace"], l["type"]] for l in ctx["financing"]["loans"]]
        add_generic_table(doc, ["Loan", "Amount", "Rate", "Term (m)", "Grace (m)", "Repayment"], rows, numeric_cols=(1,))
    if ctx["financing"]["grants"]:
        add_subheading(doc, "Grants")
        rows = [[g["name"], g["amount"], g["date"], g["restrictions"]] for g in ctx["financing"]["grants"]]
        add_generic_table(doc, ["Grant", "Amount", "Expected date", "Restrictions"], rows, numeric_cols=(1,))
    add_subheading(doc, "Tax and VAT")
    add_kpi_table(doc, ctx["tax"])
    add_subheading(doc, "KPI targets")
    add_kpi_table(doc, ctx["kpis"])


class _HtmlToDocx(HTMLParser):
    """Minimal HTML -> python-docx renderer for TipTap rich text.

    Preserves paragraphs, H2/H3 headings, bold/italic/underline/strikethrough,
    bullet + numbered lists, blockquotes, links (as text), <br> and <hr>.
    Anything unknown degrades gracefully to a styled paragraph.
    """

    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.para = None
        self.fmt = {"bold": 0, "italic": 0, "underline": 0, "strike": 0}
        self.list_stack: list[str] = []

    def _flush(self):
        self.para = None

    def handle_starttag(self, tag, attrs):
        if tag == "p":
            self.para = self.doc.add_paragraph()
        elif tag in ("h1", "h2"):
            self.para = self.doc.add_heading("", level=3)
        elif tag in ("h3", "h4", "h5"):
            self.para = self.doc.add_heading("", level=4)
        elif tag == "blockquote":
            self.para = self.doc.add_paragraph()
            self.para.paragraph_format.left_indent = Inches(0.3)
        elif tag in ("strong", "b"):
            self.fmt["bold"] += 1
        elif tag in ("em", "i"):
            self.fmt["italic"] += 1
        elif tag == "u":
            self.fmt["underline"] += 1
        elif tag in ("s", "strike", "del"):
            self.fmt["strike"] += 1
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
        elif tag == "li":
            style = "List Bullet" if (self.list_stack and self.list_stack[-1] == "ul") else "List Number"
            try:
                self.para = self.doc.add_paragraph(style=style)
            except Exception:
                self.para = self.doc.add_paragraph()
        elif tag == "br":
            if self.para is not None:
                self.para.add_run().add_break()
        elif tag == "hr":
            self.doc.add_paragraph("―" * 24)
            self._flush()

    def handle_endtag(self, tag):
        if tag in ("strong", "b"):
            self.fmt["bold"] = max(0, self.fmt["bold"] - 1)
        elif tag in ("em", "i"):
            self.fmt["italic"] = max(0, self.fmt["italic"] - 1)
        elif tag == "u":
            self.fmt["underline"] = max(0, self.fmt["underline"] - 1)
        elif tag in ("s", "strike", "del"):
            self.fmt["strike"] = max(0, self.fmt["strike"] - 1)
        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
        elif tag in ("p", "h1", "h2", "h3", "h4", "h5", "li", "blockquote"):
            self._flush()

    def handle_data(self, data):
        text = re.sub(r"[ \t\r\n]+", " ", data)
        if not text or (not text.strip() and self.para is None):
            return
        if self.para is None:
            self.para = self.doc.add_paragraph()
        run = self.para.add_run(text)
        if self.fmt["bold"]:
            run.bold = True
        if self.fmt["italic"]:
            run.italic = True
        if self.fmt["underline"]:
            run.underline = True
        if self.fmt["strike"]:
            run.font.strike = True


def _html_to_docx(doc, html: str) -> None:
    cleaned = re.sub(r"<script.*?</script>", "", html or "", flags=re.S | re.I)
    parser = _HtmlToDocx(doc)
    try:
        parser.feed(cleaned)
        parser.close()
    except Exception:
        add_paragraph_block(doc, re.sub(r"<[^>]+>", " ", html or ""))


def _add_topic_images(doc, topic):
    for img in topic.get("images", []):
        try:
            p = doc.add_paragraph()
            align = (img.get("alignment") or "center").lower()
            p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                           "full_width": WD_ALIGN_PARAGRAPH.CENTER}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            pct = img.get("width_pct") or 100
            if align == "full_width":
                pct = 100
            width = Inches(6.3 * min(max(int(pct), 10), 100) / 100.0)
            p.add_run().add_picture(img["path"], width=width)
            if img.get("caption"):
                cap = add_paragraph_block(doc, img["caption"], italic=True, size=8.5, color=SLATE)
                cap.alignment = p.alignment
        except Exception:
            continue


def add_textual_business_plan(doc, ctx):
    tp = ctx.get("text_plan") or {}
    if not tp.get("has_content"):
        return
    for sec in tp["sections"]:
        add_heading(doc, sec["title"], level=1)
        if sec.get("description"):
            add_paragraph_block(doc, sec["description"], italic=True, size=9.5, color=SLATE)
        for topic in sec["topics"]:
            add_subheading(doc, topic["title"])
            if topic.get("guidance"):
                add_warning_box(doc, f"Guidance: {topic['guidance']}", "info")
            content = (topic.get("content_html") or "").strip()
            if content:
                _html_to_docx(doc, content)
            elif (topic.get("plain_text") or "").strip():
                add_paragraph_block(doc, topic["plain_text"])
            _add_topic_images(doc, topic)


def add_financial_statements(doc, ctx):
    # The three statements are wide, so they print on landscape pages; the rest
    # of the report stays portrait (a portrait section is restored in section 9).
    _start_section(doc, landscape=True)
    add_heading(doc, "8. Financial Statements", level=1, page_break=False)
    add_paragraph_block(doc, f"All figures in {ctx['currency']}. Negative amounts are shown in parentheses; "
                             f"a dash indicates a nil balance.", italic=True, size=9)
    add_subheading(doc, "Statement of Profit or Loss")
    add_financial_table(doc, ctx["periods"], ctx["income_statement"]["rows"], ctx["currency"])
    add_subheading(doc, "Statement of Financial Position")
    add_financial_table(doc, ctx["periods"], ctx["balance_sheet"]["rows"], ctx["currency"])
    add_subheading(doc, "Statement of Cash Flows (indirect method)")
    add_financial_table(doc, ctx["periods"], ctx["cash_flow"]["rows"], ctx["currency"])


def _add_chart_image(doc, png, insight=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png), width=Inches(6.3))
    if insight:
        add_paragraph_block(doc, insight, italic=True, size=9, color=SLATE)


def add_financial_analysis(doc, ctx):
    # Return to portrait after the landscape financial statements (section 8).
    _start_section(doc, landscape=False)
    add_heading(doc, "9. Financial Analysis", level=1, page_break=False)
    add_subheading(doc, "Key performance indicators")
    add_kpi_table(doc, [{"label": k["label"], "value": k["value"]} for k in ctx["fa_kpis"]])

    if ctx["options"].include_charts and ctx.get("chart_sections"):
        add_paragraph_block(doc, "The charts below are generated from the projected financial statements and mirror "
                                 "the application's Financial Analysis dashboard.", italic=True, size=9, color=SLATE)
        for sec in ctx["chart_sections"]:
            add_subheading(doc, sec["title"])
            if sec.get("description"):
                add_paragraph_block(doc, sec["description"], size=9.5, color=SLATE)
            for rendered in charts.render_report_charts(sec["charts"], ctx["currency"]):
                _add_chart_image(doc, rendered["png"], rendered.get("insight"))


def add_scenario_comparison(doc, ctx):
    add_heading(doc, "10. Scenario Comparison", level=1)
    if ctx["options"].include_charts and ctx.get("scenario_metrics"):
        for rendered in charts.render_scenario_charts(ctx["scenario_metrics"], ctx["scenario_periods"], ctx["currency"]):
            _add_chart_image(doc, rendered["png"])
        add_subheading(doc, "Scenario summary")
    sc = ctx["scenario_comparison"]
    table = doc.add_table(rows=1, cols=4)
    _no_borders_table(table)
    for i, h in enumerate(["Metric", "Base Case", "Conservative Case", "Optimistic Case"]):
        _hdr_cell(table.rows[0].cells[i], h, left=(i == 0))
    for row in sc["rows"]:
        cells = table.add_row().cells
        _label_cell(cells[0], row["label"])
        for i, scen in enumerate(("base", "conservative", "optimistic")):
            v = row["values"].get(scen)
            _num_cell(cells[1 + i], rd.fmt_pct(v) if row["format"] == "percent" else rd.fmt_num(v))


def add_risks_and_warnings(doc, ctx):
    if not ctx["options"].include_warnings:
        return
    add_heading(doc, "11. Risks, Warnings and Reconciliations", level=1)
    add_subheading(doc, "Reconciliation checks")
    for rec in ctx["reconciliations"]:
        add_warning_box(doc, f"{rec['label']}: {'PASS' if rec['status'] else 'REVIEW REQUIRED'}",
                        "info" if rec["status"] else "critical")
    add_subheading(doc, "Warnings")
    if not ctx["warnings"]:
        add_paragraph_block(doc, "No warnings — the projection is internally consistent.", italic=True)
    for w in ctx["warnings"]:
        add_warning_box(doc, w["message"], w["severity"])


def add_appendices(doc, ctx):
    if not ctx["options"].include_appendices:
        return
    add_heading(doc, "12. Appendices", level=1)
    add_subheading(doc, "Appendix A — Revenue by stream")
    add_financial_table(doc, ctx["periods"], _cat_rows(ctx["revenue_table"], "Total revenue"), ctx["currency"], total_col=True)
    add_subheading(doc, "Appendix E — Fixed asset register")
    rows = [[a["name"], a["category"], a["purchase_amount"], a["useful_life"], a["residual"], a["nbv"]] for a in ctx["fixed_assets"]]
    add_generic_table(doc, ["Asset", "Category", "Cost", "Life", "Residual", "NBV"], rows, numeric_cols=(2, 4, 5))
    add_subheading(doc, "Appendix F — Loan schedule")
    if ctx["financing"]["loans"]:
        rows = [[l["name"], l["lender"], l["amount"], rd.fmt_pct(l["rate"]), l["term"], l["type"]] for l in ctx["financing"]["loans"]]
        add_generic_table(doc, ["Loan", "Lender", "Amount", "Rate", "Term", "Repayment"], rows, numeric_cols=(2,))
    add_subheading(doc, "Appendix I — Reconciliation checks")
    for rec in ctx["reconciliations"]:
        add_paragraph_block(doc, f"• {rec['label']}: {'PASS' if rec['status'] else 'REVIEW'}", size=9)


def _cat_rows(cat_table, total_label):
    rows = [{"label": r["label"], "values": r["values"], "kind": "line"} for r in cat_table]
    if cat_table:
        n = len(cat_table[0]["values"])
        totals = [sum(r["values"][i] for r in cat_table) for i in range(n)]
        rows.append({"label": total_label, "values": totals, "kind": "subtotal"})
    return rows


# --------------------------------------------------------------------------
# entrypoint
# --------------------------------------------------------------------------
def generate_business_plan_docx(project: BusinessPlanProject, request) -> ReportFile:
    ctx = rd.build_report_context(project, request.scenario, request.view, request)
    doc = Document()
    set_document_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    _header_footer(doc, ctx)

    add_cover_page(doc, ctx)
    add_table_of_contents_placeholder(doc)
    add_executive_summary(doc, ctx)
    add_business_overview(doc, ctx)
    add_products_revenue(doc, ctx)
    add_assumptions_sections(doc, ctx)
    add_financial_statements(doc, ctx)
    add_financial_analysis(doc, ctx)
    add_scenario_comparison(doc, ctx)
    add_risks_and_warnings(doc, ctx)
    add_appendices(doc, ctx)

    report_id = uuid.uuid4().hex
    name = rd.sanitize_filename(f"business_plan_report_{ctx['meta']['company']}_{request.scenario}_{datetime.now():%Y%m%d}")
    file_name = f"{name}_{report_id[:6]}.docx"
    path = rd.report_path(project.id, file_name)
    doc.save(str(path))
    size = path.stat().st_size

    entry = ReportFile(report_id=report_id, file_name=file_name, format="docx", scenario=request.scenario,
                       view=request.view, report_style=request.report_style, status="ready",
                       size_bytes=size, created_at=datetime.now(timezone.utc))
    rd.add_index_entry(project.id, entry.model_dump(mode="json"))
    return entry
